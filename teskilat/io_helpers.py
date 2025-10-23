# dosya okuma/yazma (PDF/DOCX/TXT)

# -*- coding: utf-8 -*-
# okuma-yazma yardımcıları; kütüphane varsa kullanırım, yoksa zarifçe düşer

import os
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import docx as docx_lib
except Exception:
    docx_lib = None

try:
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4
except Exception:
    rl_canvas = None
    A4 = None

def read_text_from_path(path):
    name, ext = os.path.splitext(path); ext = ext.lower()
    try:
        if ext in ('.txt', '.md', '.rtf'):
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        if ext == '.pdf' and PdfReader is not None:
            txt = []
            reader = PdfReader(path)
            for page in reader.pages:
                try:
                    txt.append(page.extract_text() or '')
                except Exception:
                    txt.append('')
            if not any(x.strip() for x in txt):
                return None
            return "\n".join(txt)
        if ext in ('.docx', '.doc') and docx_lib is not None:
            try:
                d = docx_lib.Document(path)
                return "\n".join(p.text for p in d.paragraphs)
            except Exception:
                return None
    except Exception:
        return None
    return None

def write_text_same_type(dest, text):
    # aynı uzantıyla yazmaya çalış; olmadı .txt’ye düş
    name, ext = os.path.splitext(dest); ext = ext.lower()
    try:
        if ext in ('.txt', '.md', '.rtf'):
            with open(dest, 'w', encoding='utf-8') as f:
                f.write(text)
            return dest, False
        if ext == '.pdf' and rl_canvas is not None and A4 is not None:
            from reportlab.pdfgen.canvas import Canvas
            c = Canvas(dest, pagesize=A4)
            w, h = A4; x, y = 50, h - 60
            import textwrap
            for para in text.split('\n'):
                for line in textwrap.wrap(para, width=95):
                    c.drawString(x, y, line); y -= 16
                    if y < 60: c.showPage(); y = h - 60
            c.save()
            return dest, False
        if ext in ('.docx', '.doc') and docx_lib is not None:
            try:
                d = docx_lib.Document()
                for para in text.split('\n'):
                    d.add_paragraph(para)
                out = dest if ext == '.docx' else (name + '.docx')
                d.save(out)
                return out, False
            except Exception:
                pass
        # fallback
        alt = name + '.txt'
        with open(alt, 'w', encoding='utf-8') as f:
            f.write(text)
        return alt, True
    except Exception:
        alt = name + '.txt'
        try:
            with open(alt, 'w', encoding='utf-8') as f:
                f.write(text)
            return alt, True
        except Exception:
            return dest, True